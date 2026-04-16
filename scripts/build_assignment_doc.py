"""
Собирает ответ на задание в DOCX и пытается экспортировать PDF через docx2pdf (нужен Microsoft Word на Windows).

Запуск из корня репозитория:
  pip install -r requirements-docs.txt
  python scripts/build_assignment_doc.py
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt


def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)


def _add_paragraphs(doc: Document, *lines: str) -> None:
    for line in lines:
        doc.add_paragraph(line)


def build_document() -> Document:
    doc = Document()
    _set_default_font(doc)

    doc.add_heading("Практическое задание: ИИ-ассистент и сценарий «сбор лидов»", 0)
    _add_paragraphs(
        doc,
        "Дисциплина: кейс по входящим заявкам / диалоговому ИИ-ассистенту.",
        "Формат ответа: анализ исходной схемы курса, проектирование нового бизнес-сценария, описание workflow, краткий отчёт.",
        "Репозиторий проекта приведён к реализации сценария квалификации лидов для отдела продаж (см. код).",
    )

    doc.add_heading("1. Исходная структура курса (техподдержка)", level=1)
    _add_paragraphs(
        doc,
        "Модель SupportTicket содержала поля: имя, контакт, краткое описание проблемы, время возникновения, место проявления, приоритет (литерал из трёх значений).",
        "В Pydantic все поля были опциональны, но метод is_complete() требовал заполнения всех шести полей для отправки заявки.",
        "SupportSession хранила user_id, chat_id, данные профиля Telegram, флаги started/submitted, объект заявки и историю диалога (до 20 сообщений).",
        "AssistantTurn включал reply (текст пользователю), extracted_ticket (частично заполненная заявка) и ready_to_submit.",
        "Нейросеть вызывалась через Chat Completions с response_format json_schema: в теле ответа приходила JSON-строка с полями reply, extracted_ticket, ready_to_submit, затем выполнялся json.loads и валидация Pydantic.",
        "Workflow объединял extracted_ticket в сессию и отправлял сообщение оператору только при is_complete() и ready_to_submit=true.",
    )

    doc.add_heading("2. Выбранный бизнес-сценарий", level=1)
    _add_paragraphs(
        doc,
        "Сценарий: сбор и первичная квалификация лидов для отдела продаж из Telegram.",
        "Цель — получить контакт, контекст компании, формулировку потребности, горизонт решения и оценку «температуры» лида, а не диагностику инцидента.",
    )

    doc.add_heading("3. Спроектированные изменения схемы и промпта", level=1)
    _add_paragraphs(
        doc,
        "Заявка переименована в смысле полей на SalesLead: name, contact, company, need_summary, timeline, lead_temperature; опционально lead_source (откуда узнали о компании).",
        "Критерии полноты заявки зафиксированы константой REQUIRED_FIELDS_FOR_COMPLETION и методом missing_required_fields(); контакт должен проходить проверку (телефон 10–15 цифр, email или @username).",
        "Удалены поля occurred_at, location и приоритет ТП; вместо problem_summary используется need_summary; timeline и lead_temperature заданы перечислениями (Literal / enum в JSON Schema).",
        "Сессия LeadSession хранит lead: SalesLead; AssistantTurn возвращает extracted_lead.",
        "Системный промпт переведён на роль пресейла: перечень полей, явный блок completion_criteria, правила не выдумывать данные, уточнение при ответе вне enum, порядок сбора данных, условие ready_to_submit.",
        "JSON Schema (strict) синхронизирована с Pydantic-моделью; в user-промпт передаются missing_required_fields_for_completion; при сбое API сохранён локальный fallback в OpenAILeadAssistant.",
        "Шаблон уведомления оператора заменён на блок «НОВЫЙ ЛИД (ПРОДАЖИ)» с каналом Telegram, источником лида, полями заявки и идентификаторами Telegram; опционально POST на CRM_WEBHOOK_URL после уведомления в чат с заголовком Idempotency-Key.",
        "Метрики качества диалога: счётчики ходов / преждевременных ready_to_submit / успешных отправок и логи lead_quality_*; A/B системного промпта через PROMPT_VARIANT=default|alt; извлечение контакта вынесено в services/lead_contact_tools.py.",
    )

    doc.add_heading("4. Логика workflow бота", level=1)
    _add_paragraphs(
        doc,
        "Шаг 1. Пользователь запускает /start — приветствие в стиле продаж, при наличии @username контакт может быть предзаполнен.",
        "Шаг 2. Сбор имени и контакта при отсутствии в lead.",
        "Шаг 3. Уточнение компании или типа клиента (B2B / частное лицо / ИП).",
        "Шаг 4. Краткое описание интереса (продукт, услуга, задача), при уточнениях обновление need_summary.",
        "Шаг 5. Сроки принятия решения — выбор из фиксированных значений timeline; при некорректном ответе повтор вопроса с перечислением вариантов.",
        "Шаг 6. Температура лида (горячий / тёплый / холодный) по смыслу ответа.",
        "Шаг 6а. По желанию — откуда узнали о компании (lead_source).",
        "Шаг 7. При полной заявке (is_complete) и ready_to_submit=true — отправка в чат менеджеров, при необходимости вебхук в CRM, финальное сообщение пользователю; если модель вернула ready_to_submit при неполных данных — уточняющий ответ без отправки.",
    )

    doc.add_heading("5. Краткий отчёт с обоснованием правок", level=1)
    _add_paragraphs(
        doc,
        "Исходный контур техподдержки хорошо разделял пользовательский текст и структурированные поля через JSON Schema, поэтому транспорт (Telegram + OpenAI) оставлен тем же, изменена только предметная модель данных.",
        "Поля «когда сломалось» и «где проявляется» для продаж малополезны, поэтому заменены на company и need_summary, а приоритет инцидента — на lead_temperature и timeline, что ближе к языку CRM и отдела продаж.",
        "Объединение extracted_lead в сессию через merge и условие отправки по is_complete() и ready_to_submit сохраняет предсказуемость и защиту от преждевременной отправки неполного лида.",
        "Промпт явно запрещает выдумывать компанию и бюджет и требует согласования enum-значений, что снижает риск невалидного JSON при strict-схеме.",
        "Fallback-логика переписана под новые вопросы, чтобы при недоступности API бот оставался работоспособным на демонстрации и учебном тестировании.",
        "Уведомление в канал переформатировано под лид, чтобы менеджеру не приходилось ментально маппить поля ТП на продажи.",
        "Документ DOCX и PDF формируются скриптом scripts/build_assignment_doc.py для прикрепления к сдаче задания вместе с ссылкой на репозиторий.",
    )

    doc.add_heading("Приложение: генерация файлов", level=1)
    _add_paragraphs(
        doc,
        "DOCX создаётся библиотекой python-docx в каталоге assignment/ в корне проекта.",
        "PDF: при установленном Microsoft Word на Windows команда docx2pdf конвертирует DOCX в PDF; при ошибке остаётся только DOCX, его можно вручную «Сохранить как PDF» из Word.",
        "Команды: pip install -r requirements-docs.txt && python scripts/build_assignment_doc.py",
    )

    return doc


def try_export_pdf(docx_path: Path, pdf_path: Path) -> bool:
    try:
        from docx2pdf import convert

        convert(str(docx_path), str(pdf_path))
    except Exception as exc:
        print(f"Не удалось получить PDF автоматически ({exc}). Сохраните DOCX в PDF из Word вручную.")
        return False
    return pdf_path.is_file()


def main() -> int:
    root = Path(__file__).resolve().parent.parent
    out_dir = root / "assignment"
    out_dir.mkdir(parents=True, exist_ok=True)

    docx_path = out_dir / "Ответ_на_задание.docx"
    pdf_path = out_dir / "Ответ_на_задание.pdf"

    doc = build_document()
    doc.save(docx_path)
    print(f"Записан DOCX: {docx_path}")

    if try_export_pdf(docx_path, pdf_path):
        print(f"Записан PDF: {pdf_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
